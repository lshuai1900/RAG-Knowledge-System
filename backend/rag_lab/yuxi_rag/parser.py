from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from .loader import DocumentFile
except ImportError:  # pragma: no cover - direct script fallback
    from loader import DocumentFile


@dataclass(slots=True)
class ParsedParagraph:
    text: str
    metadata: dict[str, Any]


_PARAGRAPH_SPLIT = re.compile(r"\n\s*\n+")


def _split_paragraphs(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [p.strip() for p in _PARAGRAPH_SPLIT.split(normalized) if p.strip()]
    if paragraphs:
        return paragraphs
    return [line.strip() for line in normalized.split("\n") if line.strip()]


def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def _parse_text_like(doc: DocumentFile) -> list[ParsedParagraph]:
    text = _read_text(doc.path)
    return [
        ParsedParagraph(
            text=paragraph,
            metadata={
                "source": doc.source,
                "file_id": doc.file_id,
                "extension": doc.extension,
                "paragraph": idx,
            },
        )
        for idx, paragraph in enumerate(_split_paragraphs(text))
    ]


def _parse_pdf(doc: DocumentFile) -> list[ParsedParagraph]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency issue
        raise RuntimeError("pypdf is required to parse PDF files") from exc

    reader = PdfReader(str(doc.path))
    paragraphs: list[ParsedParagraph] = []
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        for paragraph_idx, paragraph in enumerate(_split_paragraphs(text)):
            paragraphs.append(
                ParsedParagraph(
                    text=paragraph,
                    metadata={
                        "source": doc.source,
                        "file_id": doc.file_id,
                        "extension": doc.extension,
                        "page": page_idx,
                        "paragraph": paragraph_idx,
                    },
                )
            )
    return paragraphs


def _parse_docx(doc: DocumentFile) -> list[ParsedParagraph]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency issue
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    document = Document(str(doc.path))
    blocks: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        header = rows[0]
        blocks.append(f"| {' | '.join(header)} |")
        blocks.append(f"| {' | '.join(['---'] * len(header))} |")
        for row in rows[1:]:
            normalized = row + [""] * (len(header) - len(row))
            blocks.append(f"| {' | '.join(normalized[:len(header)])} |")

    return [
        ParsedParagraph(
            text=paragraph,
            metadata={
                "source": doc.source,
                "file_id": doc.file_id,
                "extension": doc.extension,
                "paragraph": idx,
            },
        )
        for idx, paragraph in enumerate(blocks)
    ]


def parse_document(doc: DocumentFile) -> list[ParsedParagraph]:
    if doc.extension in {".txt", ".md"}:
        return _parse_text_like(doc)
    if doc.extension == ".pdf":
        return _parse_pdf(doc)
    if doc.extension == ".docx":
        return _parse_docx(doc)
    raise ValueError(f"Unsupported file type: {doc.extension}")


def parse_documents(docs: list[DocumentFile]) -> list[ParsedParagraph]:
    paragraphs: list[ParsedParagraph] = []
    for doc in docs:
        paragraphs.extend(parse_document(doc))
    return paragraphs
