from __future__ import annotations

from pathlib import Path

from .base import BaseParser, ParseResult, format_parse_result


class DocxParser(BaseParser):
    name = "docx"

    @classmethod
    def accepts(cls) -> set[str]:
        return {".docx", ".doc"}

    def parse(self, path: Path) -> ParseResult:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse DOCX files") from exc

        document = Document(str(path))
        blocks: list[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                blocks.append(text)

        for table in document.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            header = rows[0]
            blocks.append(f"| {' | '.join(header)} |")
            blocks.append(f"| {' | '.join(['---'] * len(header))} |")
            for row in rows[1:]:
                normalized = row + [""] * (len(header) - len(row))
                blocks.append(f"| {' | '.join(normalized[:len(header)])} |")

        return format_parse_result(
            self.name,
            "\n\n".join(blocks),
            filename=path.name,
            file_type="docx",
            paragraph_count=len(blocks),
        )
