import os
import re
import asyncio
import logging
import sys
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.config import settings

logger = logging.getLogger(__name__)

_SUPPORTED_STRATEGIES = {"recursive", "semantic", "paragraph",
                         "markdown_header", "sentence_window"}
_SENTENCE_SPLIT = re.compile(r"(?<=[。！？!?；;\.])\s*")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def _resolve_strategy() -> str:
    raw = (os.getenv("CHUNK_STRATEGY") or settings.CHUNK_STRATEGY or "recursive").strip().lower()
    if raw == "semantic":
        return "paragraph"
    if raw not in _SUPPORTED_STRATEGIES:
        logger.warning("[DocumentService] Unknown CHUNK_STRATEGY=%s; falling back to recursive", raw)
        return "recursive"
    return raw


def _parse_file(path: str, ext: str) -> str:
    """Parse a single file using the new parser registry."""
    try:
        backend_dir = Path(__file__).resolve().parents[1]
        rag_lab_dir = backend_dir / "rag_lab"
        if str(rag_lab_dir) not in sys.path:
            sys.path.insert(0, str(rag_lab_dir))
        from yuxi_rag.parsers.registry import parse_file as parse_with_registry
        result = parse_with_registry(Path(path))
        return result.text
    except Exception:
        logger.debug("[DocumentService] Parser registry unavailable; using built-in loader for %s", path)
        return _fallback_load(path, ext)


def _fallback_load(path: str, ext: str) -> str:
    """Fallback loader when parser registry is not available."""
    ext = ext.lower()
    if not ext.startswith("."):
        ext = f".{ext}"

    if ext in {".txt", ".md", ".markdown"}:
        for enc in ("utf-8", "utf-8-sig", "gb18030"):
            try:
                return Path(path).read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return Path(path).read_text(encoding="utf-8", errors="ignore")

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf is required to parse PDF files") from exc
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)

    if ext in {".docx", ".doc"}:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is required to parse DOCX files") from exc
        doc = DocxDocument(path)
        blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(blocks)

    # Last resort: LangChain loaders
    from langchain_community.document_loaders import TextLoader
    loader = TextLoader(path, encoding="utf-8")
    docs = loader.load()
    return "\n\n".join(d.page_content for d in docs)


class DocumentService:
    """Document loading + chunking service.

    Uses the new parser registry when available, with LangChain fallback
    for unsupported formats.  Supports four chunk strategies:
    recursive / paragraph / markdown_header / sentence_window.
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = settings.MIN_CHUNK_SIZE
        self.max_chunk_size = settings.MAX_CHUNK_SIZE
        self._recursive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ". ", " ", ""],
        )

    async def load_and_chunk(
        self, file_path: str, file_extension: str,
        doc_id: str = "", filename: str = "",
    ) -> list[Document]:
        ext = file_extension.lower()
        if not ext.startswith("."):
            ext = f".{ext}"

        strategy = _resolve_strategy()

        # Parse document using new parser registry (with fallback)
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, _parse_file, file_path, ext)

        if not text or not text.strip():
            raise ValueError(f"Document produced no text content: {filename}")

        # Chunk according to strategy
        chunks = await self._chunk(text, strategy, filename, doc_id)

        if not chunks:
            raise ValueError(f"Document produced no chunks: {filename}")

        return chunks

    async def _chunk(
        self, text: str, strategy: str, filename: str, doc_id: str,
    ) -> list[Document]:
        loop = asyncio.get_event_loop()

        if strategy == "markdown_header":
            return await loop.run_in_executor(
                None, self._chunk_markdown_header, text, filename, doc_id,
            )
        elif strategy == "sentence_window":
            return await loop.run_in_executor(
                None, self._chunk_sentence_window, text, filename, doc_id,
            )
        elif strategy == "paragraph":
            return await loop.run_in_executor(
                None, self._chunk_paragraph, text, filename, doc_id,
            )
        else:
            # recursive (default)
            return await loop.run_in_executor(
                None, self._chunk_recursive, text, filename, doc_id,
            )

    # ── Chunk strategies ──────────────────────────────────────────────

    def _chunk_recursive(
        self, text: str, filename: str, doc_id: str,
    ) -> list[Document]:
        doc = Document(page_content=text, metadata={"source": filename})
        chunks = self._recursive_splitter.split_documents([doc])
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "chunk_index": i,
                "chunk_id": f"{doc_id}_{i}",
                "chunk_strategy": "recursive",
                "document_id": doc_id,
                "filename": filename,
            })
        return chunks

    def _chunk_paragraph(
        self, text: str, filename: str, doc_id: str,
    ) -> list[Document]:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        if not paragraphs:
            paragraphs = [text]

        chunks: list[Document] = []
        idx = 0
        for para in paragraphs:
            if len(para) <= self.chunk_size:
                chunks.append(Document(
                    page_content=para,
                    metadata={
                        "source": filename,
                        "chunk_index": idx,
                        "chunk_id": f"{doc_id}_{idx}",
                        "chunk_strategy": "paragraph",
                        "document_id": doc_id,
                        "filename": filename,
                        "char_count": len(para),
                    },
                ))
                idx += 1
            else:
                sub_chunks = self._split_long_text(para, filename, doc_id, idx, "paragraph")
                chunks.extend(sub_chunks)
                idx += len(sub_chunks)
        return chunks

    def _chunk_markdown_header(
        self, text: str, filename: str, doc_id: str,
    ) -> list[Document]:
        headings = list(_HEADING_RE.finditer(text))

        if not headings:
            # No headings found — fall back to paragraph chunking
            logger.info("[DocumentService] No markdown headings found in %s; using paragraph strategy", filename)
            return self._chunk_paragraph(text, filename, doc_id)

        sections: list[tuple[str, str | None, str]] = []
        for i, match in enumerate(headings):
            level = len(match.group(1))
            title = match.group(2).strip()
            start = match.end()
            end = headings[i + 1].start() if i + 1 < len(headings) else len(text)
            body = text[start:end].strip()
            section_path = title
            sections.append((body, section_path, title))

        # Also capture content before first heading
        if headings and headings[0].start() > 0:
            preamble = text[:headings[0].start()].strip()
            if preamble:
                sections.insert(0, (preamble, None, None))

        chunks: list[Document] = []
        idx = 0
        for body, section_path, section_title in sections:
            if not body:
                continue
            if len(body) <= self.chunk_size:
                chunks.append(Document(
                    page_content=body,
                    metadata={
                        "source": filename,
                        "chunk_index": idx,
                        "chunk_id": f"{doc_id}_{idx}",
                        "chunk_strategy": "markdown_header",
                        "document_id": doc_id,
                        "filename": filename,
                        "section_title": section_title or "",
                        "section_path": section_path or "",
                        "char_count": len(body),
                    },
                ))
                idx += 1
            else:
                sub_chunks = self._split_long_text(body, filename, doc_id, idx, "markdown_header")
                for sc in sub_chunks:
                    sc.metadata["section_title"] = section_title or ""
                    sc.metadata["section_path"] = section_path or ""
                chunks.extend(sub_chunks)
                idx += len(sub_chunks)
        return chunks

    def _chunk_sentence_window(
        self, text: str, filename: str, doc_id: str,
    ) -> list[Document]:
        sentences = [s.strip() for s in _SENTENCE_SPLIT.split(text) if s.strip()]
        if not sentences:
            sentences = [text]

        chunks: list[Document] = []
        idx = 0
        i = 0
        while i < len(sentences):
            window = sentences[i]
            j = i + 1
            while j < len(sentences) and len(window) + len(sentences[j]) + 1 <= self.chunk_size:
                window += " " + sentences[j]
                j += 1

            chunks.append(Document(
                page_content=window,
                metadata={
                    "source": filename,
                    "chunk_index": idx,
                    "chunk_id": f"{doc_id}_{idx}",
                    "chunk_strategy": "sentence_window",
                    "document_id": doc_id,
                    "filename": filename,
                    "char_count": len(window),
                    "sentence_start": i,
                    "sentence_end": j - 1,
                },
            ))
            idx += 1

            if j == i:
                i += 1
            else:
                overlap_steps = max(1, self.chunk_overlap // 50)
                i = max(i + 1, j - overlap_steps)

        return chunks

    def _split_long_text(
        self, text: str, filename: str, doc_id: str,
        start_idx: int, strategy: str,
    ) -> list[Document]:
        """Split long text into overlapping chunks using recursive splitter."""
        doc = Document(page_content=text, metadata={"source": filename})
        sub_chunks = self._recursive_splitter.split_documents([doc])
        results: list[Document] = []
        for offset, chunk in enumerate(sub_chunks):
            chunk.metadata.update({
                "chunk_index": start_idx + offset,
                "chunk_id": f"{doc_id}_{start_idx + offset}",
                "chunk_strategy": strategy,
                "document_id": doc_id,
                "filename": filename,
                "char_count": len(chunk.page_content),
            })
            results.append(chunk)
        return results
